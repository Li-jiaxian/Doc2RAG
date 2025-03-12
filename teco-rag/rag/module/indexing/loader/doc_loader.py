# BSD 3- Clause License Copyright (c) 2023, Tecorigin Co., Ltd. All rights
# reserved.
# Redistribution and use in source and binary forms, with or without
# modification, are permitted provided that the following conditions are met:
# Redistributions of source code must retain the above copyright notice,
# this list of conditions and the following disclaimer.
# Redistributions in binary form must reproduce the above copyright notice,
# this list of conditions and the following disclaimer in the documentation
# and/or other materials provided with the distribution.
# Neither the name of the copyright holder nor the names of its contributors
# may be used to endorse or promote products derived from this software
# without specific prior written permission.
#
# THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
# AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
# IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE
# ARE DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE
# LIABLE FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR
# CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF
# SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS
# INTERRUPTION)
# HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT,
# STRICT LIABILITY,OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE)  ARISING IN ANY
# WAY OUT OF THE USE OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY
# OF SUCH DAMAGE.

from io import BytesIO
from typing import List

import numpy as np
import tqdm
from docx import Document, ImagePart
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from langchain_community.document_loaders import UnstructuredFileLoader
from PIL import Image

from rag.module.indexing.loader.utils.ocr import get_rapid_ocr


class CustomizedOcrDocLoader(UnstructuredFileLoader):
    """
    自定义的OCR文档加载器类，继承自UnstructuredFileLoader。
    这个类用于加载和处理包含文本和图像的文档，并使用OCR技术识别图像中的文本。
    """

    def _is_paragraph_end(self, text):
        """
        判断给定文本是否为段落结束。

        Args:
            text (str): 需要判断的文本。

        Returns:
            bool: 如果文本以句号或问号结尾，则返回True，否则返回False。
        """
        return text.strip()[-1] in ["。", "？"]

    def _get_elements(self) -> List:
        """
        获取文档中的所有元素（文本和图像）。

        Returns:
            List: 包含文档所有元素的列表。
        """

        def doc2text(filepath):
            """
            将文档转换为文本，包括OCR识别图像中的文本。

            Args:
                filepath (str): 文档文件路径。

            Returns:
                str: 提取的文本内容。
            """
            ocr = get_rapid_ocr()

            doc = Document(filepath)  # 无法读取doc文件, 只能读取docx文件
            resp = ""

            def iter_block_items(parent):
                """
                迭代文档中的块元素（段落和表格）。

                Args:
                    parent: 父元素（文档或单元格）。

                Yields:
                    Paragraph 或 Table: 文档中的段落或表格元素。
                """
                from docx.document import Document

                if isinstance(parent, Document):
                    parent_elm = parent.element.body
                elif isinstance(parent, _Cell):
                    parent_elm = parent._tc
                else:
                    raise ValueError("CustomizedOcrDocLoader parse fail")

                for child in parent_elm.iterchildren():
                    if isinstance(child, CT_P):
                        yield Paragraph(child, parent)
                    elif isinstance(child, CT_Tbl):
                        yield Table(child, parent)

            b_unit = tqdm.tqdm(
                total=len(doc.paragraphs) + len(doc.tables),
                desc="CustomizedOcrDocLoader block index: 0",
            )
            for i, block in enumerate(iter_block_items(doc)):
                b_unit.set_description(
                    "CustomizedOcrDocLoader block index: {}".format(i)
                )
                b_unit.refresh()
                if isinstance(block, Paragraph):
                    resp += block.text.strip() + "\n"

                    images = block._element.xpath(".//pic:pic")  # 获取所有图片
                    for image in images:
                        for img_id in image.xpath(".//a:blip/@r:embed"):  # 获取图片id
                            part = doc.part.related_parts[
                                img_id
                            ]  # 根据图片id获取对应的图片
                            if isinstance(part, ImagePart):
                                image = Image.open(BytesIO(part._blob))
                                result, _ = ocr(np.array(image))
                                if result:
                                    ocr_result = [line[1] for line in result]
                                    resp += "\n".join(ocr_result)

                elif isinstance(block, Table):
                    for row in block.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                resp += paragraph.text.strip() + "\n"

                b_unit.update(1)
            return resp

        text = doc2text(self.file_path)
        from unstructured.partition.text import partition_text

        return partition_text(text=text, **self.unstructured_kwargs)
